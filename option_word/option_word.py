from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class option_word:
    def __init__(self):
        self.document = Document()
        self.document.styles["Normal"].font.name = u"宋体"  # 设置全局字体
        self.document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.document.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)  # 设置正文全局颜色为黑色
        self.document.styles["Normal"].font.size = Pt(18)  # 设置正文全局大小为18
        self.document.styles["Heading 2"].font.size = Pt(20)  # 设置全局2级标题的字体大小为18
        self.document.styles["Heading 2"].font.color.rgb = RGBColor(0, 0, 0)  # 设置2级标题全局颜色为黑色
        self.document.styles["Heading 2"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def add_one_paragraph(self, content, style=None) -> Paragraph:
        """
        添加一个段落(正文)， 返回值是当前的段落对象
        """
        return self.document.add_paragraph(content, style)

    def save(self, doc_path):
        """
        保存文档
        """
        self.document.save(doc_path)

    @staticmethod
    def before_paragraph_insert(inserted_paragraph, content, style=None) -> Paragraph:
        """
        在某个段落前插入一个段落， 返回值是 前面的段落对象
        """
        return inserted_paragraph.insert_paragraph_before(content, style)

    def add_header(self, content, level):
        return self.document.add_heading(content, level)

    @staticmethod
    def get_style(paragraph_obj: Paragraph):
        return paragraph_obj.style

    @staticmethod
    def content_center(paragraph_obj: Paragraph):
        paragraph_obj.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


if __name__ == '__main__':
    ow = option_word()
    test = ow.add_one_paragraph("test")
    x = ow.before_paragraph_insert(test, "test1")
    ow.add_header("this is a header", level=2)
    y = ow.get_style(test)
    ow.content_center(test)
    print(y)
    ow.save("test.docx")
